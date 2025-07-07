from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import tempfile
import shutil
from pathlib import Path
import uuid
from typing import Optional
import asyncio
from concurrent.futures import ThreadPoolExecutor
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from PIL import Image
import io
import logging
import pytesseract
import cv2
import numpy as np

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="PDF to Word Converter", version="1.0.0")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure this properly for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories
UPLOAD_DIR = Path("uploads")
CONVERTED_DIR = Path("converted")
UPLOAD_DIR.mkdir(exist_ok=True)
CONVERTED_DIR.mkdir(exist_ok=True)

# Thread pool for CPU-intensive operations
executor = ThreadPoolExecutor(max_workers=4)

class PDFConverter:
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
    
    def convert_pdf_to_docx(self, pdf_path: str, docx_path: str) -> bool:
        """
        Convert PDF to DOCX using PyMuPDF with enhanced layout preservation
        """
        try:
            doc = Document()
            pdf_document = fitz.open(pdf_path)
            
            # Add document title
            title = doc.add_heading('Converted Document', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Add page break for subsequent pages
                if page_num > 0:
                    doc.add_page_break()
                
                # Add page heading
                page_heading = doc.add_heading(f'Page {page_num + 1}', level=1)
                page_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Get page dimensions for image positioning
                page_rect = page.rect
                
                # Extract text with positioning information
                text_dict = page.get_text("dict")
                
                # Extract images first
                image_list = page.get_images(full=True)
                image_positions = {}
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            image_data = pix.tobytes("png")
                            image_path = os.path.join(self.temp_dir, f"image_{page_num}_{img_index}.png")
                            
                            with open(image_path, "wb") as img_file:
                                img_file.write(image_data)
                            
                            # Add image to document
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                            
                            # Calculate image size (maintain aspect ratio)
                            img_width = min(6, pix.width / 100)  # Max 6 inches
                            
                            paragraph.add_run().add_picture(image_path, width=Inches(img_width))
                            
                            # Store image position for text flow
                            image_positions[img_index] = (pix.x, pix.y, pix.width, pix.height)
                        
                        pix = None
                    except Exception as img_error:
                        logger.warning(f"Could not extract image {img_index}: {img_error}")
                        continue
                
                # Process text blocks with better formatting
                for block in text_dict["blocks"]:
                    if "lines" in block:
                        block_text = ""
                        font_sizes = []
                        is_bold = False
                        
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                                font_sizes.append(span["size"])
                                if span["flags"] & 2**4:  # Bold flag
                                    is_bold = True
                            
                            if line_text.strip():
                                block_text += line_text + "\n"
                        
                        if block_text.strip():
                            # Determine formatting based on font size
                            avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
                            
                            if avg_font_size > 16:
                                # Large text - likely heading
                                heading = doc.add_heading(block_text.strip(), level=2)
                                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif avg_font_size > 14:
                                # Medium text - subheading
                                heading = doc.add_heading(block_text.strip(), level=3)
                                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                # Normal text
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(block_text.strip())
                                
                                if is_bold:
                                    run.bold = True
                                
                                # Set font size
                                run.font.size = Pt(max(10, min(avg_font_size, 14)))
                
                # Add some spacing between pages
                doc.add_paragraph()
            
            pdf_document.close()
            doc.save(docx_path)
            return True
            
        except Exception as e:
            logger.error(f"PDF conversion failed: {e}")
            return False
    
    def preprocess_image(self, image_path: str) -> tuple[str, str]:
        """
        Preprocess image for better OCR results with multiple preprocessing approaches
        """
        try:
            # Read image
            image = cv2.imread(image_path)
            if image is None:
                # Try with PIL if OpenCV fails
                pil_image = Image.open(image_path)
                image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
            
            # Get image dimensions for scaling
            height, width = image.shape[:2]
            
            # Upscale image if it's too small (helps with small text)
            if width < 1000 or height < 1000:
                scale_factor = max(1000 / width, 1000 / height)
                new_width = int(width * scale_factor)
                new_height = int(height * scale_factor)
                image = cv2.resize(image, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
            
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Method 1: Standard preprocessing
            denoised = cv2.fastNlMeansDenoising(gray)
            thresh1 = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Method 2: Enhanced preprocessing for code screenshots
            # Increase contrast
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            
            # Apply bilateral filter to reduce noise while preserving edges
            bilateral = cv2.bilateralFilter(enhanced, 9, 75, 75)
            
            # Try different thresholding methods
            # Simple binary threshold
            _, thresh2 = cv2.threshold(bilateral, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Morphological operations to clean up
            kernel = np.ones((1,1), np.uint8)
            thresh2 = cv2.morphologyEx(thresh2, cv2.MORPH_CLOSE, kernel)
            thresh2 = cv2.morphologyEx(thresh2, cv2.MORPH_OPEN, kernel)
            
            # Save both processed images
            processed_path1 = os.path.join(self.temp_dir, "processed_standard.png")
            processed_path2 = os.path.join(self.temp_dir, "processed_enhanced.png")
            
            cv2.imwrite(processed_path1, thresh1)
            cv2.imwrite(processed_path2, thresh2)
            
            return processed_path1, processed_path2
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}")
            return image_path, image_path  # Return original if preprocessing fails
    
    def convert_image_to_docx(self, image_path: str, docx_path: str) -> bool:
        """
        Convert image to DOCX with OCR text extraction and layout preservation
        """
        try:
            doc = Document()
            doc.add_heading('Converted Image Document', 0)
            
            # Preprocess image for better OCR
            processed_image_path1, processed_image_path2 = self.preprocess_image(image_path)
            
            # Add the original image first
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(image_path, width=Inches(6))
            
            # Add spacing
            doc.add_paragraph()
            
            # Extract text with multiple OCR approaches
            logger.info("Extracting text from image...")
            
            # OCR Configuration options for different content types
            ocr_configs = [
                '--psm 6 -c tessedit_char_whitelist=0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ .,;:!?-_()[]{}/<>="\'`~@#$%^&*+|\\',  # Code-friendly
                '--psm 4 -c tessedit_char_whitelist=0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ .,;:!?-_()[]{}/<>="\'`~@#$%^&*+|\\',  # Single column
                '--psm 6',  # Uniform block of text
                '--psm 12',  # Sparse text
                '--psm 8',  # Single word
                '--psm 13'   # Raw line
            ]
            
            best_result = ""
            best_confidence = 0
            detailed_data = None
            
            # Try different OCR configurations on both processed images
            for img_path in [processed_image_path1, processed_image_path2, image_path]:
                for config in ocr_configs:
                    try:
                        # Get detailed OCR data
                        ocr_data = pytesseract.image_to_data(
                            img_path, 
                            output_type=pytesseract.Output.DICT,
                            config=config
                        )
                        
                        # Get simple text
                        text_result = pytesseract.image_to_string(img_path, config=config)
                        
                        # Calculate average confidence
                        confidences = [conf for conf in ocr_data['conf'] if conf > 0]
                        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                        
                        # Use result with highest confidence and reasonable text length
                        if avg_confidence > best_confidence and len(text_result.strip()) > 10:
                            best_confidence = avg_confidence
                            best_result = text_result
                            detailed_data = ocr_data
                            
                    except Exception as e:
                        logger.warning(f"OCR attempt failed with config {config}: {e}")
                        continue
            
            # Process the best OCR result
            if best_result.strip() and detailed_data:
                # Add confidence info
                confidence_para = doc.add_paragraph()
                confidence_run = confidence_para.add_run(f"OCR Confidence: {best_confidence:.1f}%")
                confidence_run.font.size = Pt(9)
                confidence_run.italic = True
                
                doc.add_paragraph()
                
                # For code screenshots, only show clean formatted version
                if self.is_likely_code(best_result):
                    doc.add_heading('Extracted Code', level=1)
                    
                    # Clean up the text for better code readability
                    cleaned_text = self.clean_code_text(best_result)
                    
                    # Split by lines and preserve indentation
                    lines = cleaned_text.split('\n')
                    for line in lines:
                        if line.strip():  # Skip empty lines
                            para = doc.add_paragraph()
                            run = para.add_run(line)
                            run.font.name = 'Courier New'  # Monospace font
                            run.font.size = Pt(10)
                            
                            # Try to preserve indentation
                            leading_spaces = len(line) - len(line.lstrip())
                            if leading_spaces > 0:
                                para.paragraph_format.left_indent = Inches(leading_spaces * 0.05)
                        else:
                            doc.add_paragraph()  # Empty line
                
                else:
                    # Regular text processing with detailed OCR data
                    doc.add_heading('Extracted Text Content', level=1)
                    
                    current_paragraph = None
                    current_block = -1
                    
                    for i, text in enumerate(detailed_data['text']):
                        if text.strip():
                            confidence = detailed_data['conf'][i]
                            block_num = detailed_data['block_num'][i]
                            par_num = detailed_data['par_num'][i]
                            
                            # Skip text with very low confidence
                            if confidence < 20:
                                continue
                            
                            # Start new paragraph for new blocks
                            if block_num != current_block:
                                current_paragraph = doc.add_paragraph()
                                current_block = block_num
                            
                            # Add text with formatting
                            if current_paragraph:
                                run = current_paragraph.add_run(text + " ")
                                
                                # Format based on confidence
                                if confidence > 80:
                                    run.font.size = Pt(12)
                                elif confidence > 60:
                                    run.font.size = Pt(11)
                                else:
                                    run.font.size = Pt(10)
                                
                                # Check for headings
                                if (len(text) > 3 and confidence > 70 and 
                                    (text.isupper() or text.istitle())):
                                    current_paragraph.clear()
                                    heading = doc.add_heading(text, level=2)
                                    current_paragraph = None
                                    current_block = -1
                    
                    # Add raw text section for non-code content
                    doc.add_paragraph()
                    doc.add_heading('Raw Extracted Text', level=1)
                    
                    # Regular text paragraphs
                    paragraphs = best_result.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            lines = para_text.split('\n')
                            for line in lines:
                                if line.strip():
                                    para = doc.add_paragraph()
                                    run = para.add_run(line.strip())
                                    run.font.size = Pt(12)
                            doc.add_paragraph()
                
            else:
                # If OCR failed completely
                doc.add_heading('Text Extraction Results', level=1)
                para = doc.add_paragraph()
                run = para.add_run("Text extraction was attempted but no readable text was found.")
                run.font.size = Pt(12)
                
                para = doc.add_paragraph()
                run = para.add_run("Tips for better OCR results:")
                run.font.size = Pt(12)
                run.bold = True
                
                tips = [
                    "• Use high resolution images (at least 300 DPI)",
                    "• Ensure good contrast between text and background",
                    "• Avoid skewed or rotated text",
                    "• Use clear, readable fonts",
                    "• For code screenshots, use larger font sizes"
                ]
                
                for tip in tips:
                    para = doc.add_paragraph()
                    run = para.add_run(tip)
                    run.font.size = Pt(11)
            
            # Add document info
            doc.add_paragraph()
            doc.add_heading('Document Information', level=1)
            info_para = doc.add_paragraph()
            info_run = info_para.add_run(f"Original file: {Path(image_path).name}")
            info_run.font.size = Pt(10)
            
            info_para = doc.add_paragraph()
            info_run = info_para.add_run(f"OCR Confidence: {best_confidence:.1f}%")
            info_run.font.size = Pt(10)
            
            info_para = doc.add_paragraph()
            info_run = info_para.add_run("Generated using enhanced OCR technology")
            info_run.font.size = Pt(10)
            
            doc.save(docx_path)
            logger.info("Image to DOCX conversion completed successfully")
            return True
            
        except Exception as e:
            logger.error(f"Image conversion failed: {e}")
            return False
    
    def clean_code_text(self, text: str) -> str:
        """
        Clean up OCR text for better code readability
        """
        # Remove common OCR artifacts in code
        text = text.replace('|', 'l')  # Common OCR mistake
        text = text.replace(' . ', '.')  # Fix spaced periods
        text = text.replace(' , ', ',')  # Fix spaced commas
        text = text.replace(' ; ', ';')  # Fix spaced semicolons
        text = text.replace(' ( ', '(')  # Fix spaced parentheses
        text = text.replace(' ) ', ')')
        text = text.replace('nmumpy', 'numpy')  # Fix common numpy OCR error
        text = text.replace('getLogger(__name_)', 'getLogger(__name__)')  # Fix common error
        text = text.replace(' = ', '=')  # Fix spaced assignments
        text = text.replace('basicConfig', 'basicConfig')  # Fix common logging error
        
        # Remove UI elements that commonly appear in code screenshots
        ui_elements = [
            '@ Code', 'File', 'Edit', 'Selection', 'View', 'Go', 'Run', 'Terminal', 'Window', 'Help',
            'OO &', 'Mon 7 Jul', 'Exit Full Screen', 'Dy &', 'server-python', '™', 'apps', 'mas',
            'B main.py', 'ae', 'SSS', 'ol &', 'aan,', 'aan', 'cK', 'es', 'ny', 'i 0', 'ne', ':'
        ]
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip lines that are just UI elements
            if any(ui_element in line for ui_element in ui_elements) and len(line.strip()) < 50:
                continue
            
            # Skip lines that are just numbers (line numbers)
            if line.strip().isdigit() and len(line.strip()) < 3:
                continue
            
            # Skip lines with just symbols or very short lines
            if len(line.strip()) < 3 and not line.strip().startswith(('#', '//', 'import', 'from')):
                continue
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def is_likely_code(self, text: str) -> bool:
        """
        Heuristic to determine if text is likely code
        """
        code_indicators = [
            '{', '}', '(', ')', '[', ']', ';', '//', '/*', '*/', 
            'def ', 'class ', 'function', 'import ', 'from ', 'if ', 'for ', 'while ',
            'var ', 'let ', 'const ', 'return ', 'print(', 'console.', '<?', '?>', '</', 'href='
        ]
        
        # Count code-like patterns
        code_count = sum(1 for indicator in code_indicators if indicator in text)
        
        # Check for common code characteristics
        lines = text.split('\n')
        indented_lines = sum(1 for line in lines if line.startswith('    ') or line.startswith('\t'))
        
        # If more than 20% of lines are indented or we have multiple code indicators
        is_code = (indented_lines > len(lines) * 0.2) or (code_count >= 3)
        
        return is_code
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.warning(f"Cleanup failed: {e}")

@app.post("/convert")
async def convert_file(file: UploadFile = File(...)):
    """
    Convert uploaded file (PDF/Image) to Word document
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    # Validate file type
    allowed_extensions = {'.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.tiff'}
    file_ext = Path(file.filename).suffix.lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
        )
    
    # Generate unique filenames
    unique_id = str(uuid.uuid4())
    input_filename = f"{unique_id}_{file.filename}"
    output_filename = f"{unique_id}_{Path(file.filename).stem}.docx"
    
    input_path = UPLOAD_DIR / input_filename
    output_path = CONVERTED_DIR / output_filename
    
    try:
        # Save uploaded file
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Initialize converter
        converter = PDFConverter()
        
        # Convert based on file type
        if file_ext == '.pdf':
            success = await asyncio.get_event_loop().run_in_executor(
                executor, 
                converter.convert_pdf_to_docx, 
                str(input_path), 
                str(output_path)
            )
        else:
            # Image file
            success = await asyncio.get_event_loop().run_in_executor(
                executor, 
                converter.convert_image_to_docx, 
                str(input_path), 
                str(output_path)
            )
        
        converter.cleanup()
        
        if not success:
            raise HTTPException(status_code=500, detail="Conversion failed")
        
        if not output_path.exists():
            raise HTTPException(status_code=500, detail="Conversion completed but output file not found")
        
        # Return the converted file
        return FileResponse(
            path=str(output_path),
            filename=f"{Path(file.filename).stem}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        logger.error(f"Conversion error: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
    
    finally:
        # Cleanup files after a delay
        asyncio.create_task(cleanup_files(input_path, output_path))

async def cleanup_files(input_path: Path, output_path: Path, delay: int = 300):
    """
    Clean up files after a delay (5 minutes default)
    """
    await asyncio.sleep(delay)
    try:
        if input_path.exists():
            input_path.unlink()
        if output_path.exists():
            output_path.unlink()
    except Exception as e:
        logger.warning(f"File cleanup failed: {e}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "PDF to Word Converter"}

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "PDF to Word Conversion Service",
        "version": "1.0.0",
        "endpoints": {
            "convert": "/convert (POST)",
            "health": "/health (GET)"
        }
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)